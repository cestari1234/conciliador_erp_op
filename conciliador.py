import streamlit as st
import pandas as pd
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuração da Página ---
st.set_page_config(
    page_title="Conciliador de Vendas",
    page_icon="📊",
    layout="wide"
)

# --- Estilização Customizada (Opcional) ---
st.markdown("""
    <style>
    .main {
        padding-top: 2rem;
    }
    .stButton>button {
        width: 100%;
        margin-top: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)

# --- Funções Auxiliares ---

@st.cache_data
def load_data(file):
    """Carrega dados de CSV ou Excel."""
    try:
        if file.name.endswith('.csv'):
            return pd.read_csv(file)
        elif file.name.endswith(('.xls', '.xlsx')):
            return pd.read_excel(file)
    except Exception as e:
        st.error(f"Erro ao carregar arquivo: {e}")
        return None

def standardize_df(df, mapping, system_name_col, system_name_val):
    """Renomeia colunas e adiciona coluna de origem."""
    df_std = df.rename(columns=mapping)
    # Manter apenas as colunas mapeadas
    cols_to_keep = list(mapping.values())
    
    # Verificar se todas as colunas existem
    available_cols = [c for c in cols_to_keep if c in df_std.columns]
    df_std = df_std[available_cols]
    
    # Adicionar coluna de identificação
    df_std[system_name_col] = system_name_val
    
    return df_std

def get_operator_suggestion(name):
    """Retorna sugestão de padronização baseada em regras."""
    name_upper = str(name).upper()
    
    rules = {
        'TICKET': 'TICKET',
        'BIG': 'BIGCARD',
        'MAESTRO': 'MASTERCARD',
        'MASTER': 'MASTERCARD',
        'ALELO': 'ALELO',  # Alelo deve vir antes de Elo para evitar falso positivo
        'ELO': 'ELO',
        'VISA': 'VISA',
        'VR': 'VR',
        'SODEXO': 'PLUXEE',
        'CABAL': 'CABAL',
        'AMEX': 'AMEX',
        'COMPROCARD': 'COMPROCARD',
        'UPBRASIL': 'UPBRASIL',
        'PIX': 'PIX',
        'POLICARD': 'POLICARD',
        'NÃO INFORMADO': 'NÃO INFORMADO'
    }
    
    for key, value in rules.items():
        if key in name_upper:
            return value
    return ""

def clean_input_value(val):
    """Limpa valores de entrada (R$, %, vírgula) e converte para float."""
    if pd.isna(val) or val == "":
        return 0.0
    
    if isinstance(val, (int, float)):
        return float(val)
        
    val_str = str(val).strip()
    
    # Remover símbolos de moeda e porcentagem
    val_str = val_str.replace('R$', '').replace('%', '').replace(' ', '')
    
    # Tratar vírgula como decimal se não houver ponto, ou se houver ponto e vírgula (ex: 1.000,00)
    if ',' in val_str:
        if '.' in val_str:
            # Assumir formato brasileiro 1.000,00 -> remover ponto, trocar vírgula por ponto
            val_str = val_str.replace('.', '').replace(',', '.')
        else:
            # Assumir apenas vírgula decimal 10,00 -> trocar vírgula por ponto
            val_str = val_str.replace(',', '.')
            
    try:
        return float(val_str)
    except ValueError:
        return 0.0

def generate_word_report(erp_metrics, op_metrics, summary_df, daily_df, start_date, end_date, erp_divergences, op_divergences):
    """Gera um relatório executivo em Word (.docx)."""
    doc = Document()
    
    # --- Estilos ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título
    heading = doc.add_heading('Relatório Executivo de Conciliação', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Período
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Período: {start_date.strftime('%d/%m/%Y')} a {end_date.strftime('%d/%m/%Y')}")
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph() # Espaço
    
    # --- 1. Resumo Executivo ---
    doc.add_heading('1. Resumo Executivo', level=1)
    
    # Tabela de KPIs
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'ERP'
    hdr_cells[2].text = 'Operadora'
    
    # Preencher KPIs
    metrics = [
        ("Vendas Brutas", erp_metrics['vendas_brutas'], op_metrics['vendas_brutas']),
        ("Taxa Média", f"{erp_metrics['taxa_media']:.2f}%", f"{op_metrics['taxa_media']:.2f}%"),
        ("Total de Taxas", erp_metrics['total_taxas'], op_metrics['total_taxas']),
        ("Vendas Líquidas", erp_metrics['vendas_liquidas'], op_metrics['vendas_liquidas'])
    ]
    
    for metric, val_erp, val_op in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = val_erp if isinstance(val_erp, str) else f"R$ {val_erp:,.2f}"
        row_cells[2].text = val_op if isinstance(val_op, str) else f"R$ {val_op:,.2f}"
        
    doc.add_paragraph()
    
    # --- 2. Detalhamento Financeiro ---
    doc.add_heading('2. Detalhamento Financeiro', level=1)
    doc.add_paragraph("Abaixo o comparativo detalhado entre os valores registrados no ERP e na Operadora.")
    
    # Usar o DataFrame de resumo já calculado
    # Converter para lista de listas para tabela do Word
    
    # --- 3. Performance por Adquirente ---
    doc.add_heading('3. Performance por Adquirente', level=1)
    
    if not summary_df.empty:
        # Adicionar Tabela
        table_perf = doc.add_table(rows=1, cols=len(summary_df.columns))
        table_perf.style = 'Light Shading Accent 1'
        
        # Cabeçalho
        for i, col_name in enumerate(summary_df.columns):
            table_perf.rows[0].cells[i].text = str(col_name)
            
        # Linhas
        for index, row in summary_df.iterrows():
            row_cells = table_perf.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
    else:
        doc.add_paragraph("Nenhum dado disponível para este período.")

    # --- 4. Detalhamento - ERP -> Operadora (Divergentes) ---
    doc.add_heading('4. Detalhamento - ERP -> Operadora (Divergentes)', level=1)
    doc.add_paragraph("Registros do ERP que não foram encontrados na Operadora.")

    if not erp_divergences.empty:
        # Limitar colunas para caber na página (opcional, mas recomendado)
        # Vamos pegar as colunas principais
        cols_to_export = ['data_venda', 'valor_venda', 'operadora', 'nsu']
        # Verificar se existem
        cols_to_export = [c for c in cols_to_export if c in erp_divergences.columns]
        
        table_div_erp = doc.add_table(rows=1, cols=len(cols_to_export))
        table_div_erp.style = 'Table Grid'
        
        # Cabeçalho
        for i, col_name in enumerate(cols_to_export):
            table_div_erp.rows[0].cells[i].text = str(col_name)
            
        # Linhas (Limitar a 100 para não quebrar o Word se for muito grande, ou avisar)
        # O usuário pediu tabelas separadas, vamos colocar tudo mas com cuidado
        # Se for muito grande, o Word pode ficar lento. Vamos limitar a 500 linhas por segurança.
        max_rows = 500
        for index, row in erp_divergences.head(max_rows).iterrows():
            row_cells = table_div_erp.add_row().cells
            for i, col in enumerate(cols_to_export):
                row_cells[i].text = str(row[col])
        
        if len(erp_divergences) > max_rows:
            doc.add_paragraph(f"... e mais {len(erp_divergences) - max_rows} registros. (Exibição limitada a {max_rows} itens)")
    else:
        doc.add_paragraph("Nenhuma divergência encontrada nesta direção.")

    # --- 5. Detalhamento - Operadora -> ERP (Divergentes) ---
    doc.add_heading('5. Detalhamento - Operadora -> ERP (Divergentes)', level=1)
    doc.add_paragraph("Registros da Operadora que não foram encontrados no ERP.")

    if not op_divergences.empty:
        cols_to_export = ['data_venda', 'valor_venda', 'operadora', 'nsu']
        cols_to_export = [c for c in cols_to_export if c in op_divergences.columns]
        
        table_div_op = doc.add_table(rows=1, cols=len(cols_to_export))
        table_div_op.style = 'Table Grid'
        
        # Cabeçalho
        for i, col_name in enumerate(cols_to_export):
            table_div_op.rows[0].cells[i].text = str(col_name)
            
        # Linhas
        max_rows = 500
        for index, row in op_divergences.head(max_rows).iterrows():
            row_cells = table_div_op.add_row().cells
            for i, col in enumerate(cols_to_export):
                row_cells[i].text = str(row[col])
                
        if len(op_divergences) > max_rows:
            doc.add_paragraph(f"... e mais {len(op_divergences) - max_rows} registros. (Exibição limitada a {max_rows} itens)")
    else:
        doc.add_paragraph("Nenhuma divergência encontrada nesta direção.")

    # Rodapé
    section = doc.sections[0]
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.text = "Gerado automaticamente pelo sistema Concicard"
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Salvar em memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Interface Principal ---

st.title("📊 Conciliador de Vendas")
st.markdown("Faça o upload das planilhas do **ERP** e da **Operadora** para iniciar a conciliação.")

col1, col2 = st.columns(2)

with col1:
    st.subheader("📁 Arquivo 1: Sistema ERP")
    file_erp = st.file_uploader("Upload Planilha ERP", type=['csv', 'xlsx', 'xls'], key="erp")
    erp_name = st.selectbox("Nome do Sistema ERP", options=["Brajan", "Uniplus", "StarTwo", "Exodus", "Outro"])
    
    if erp_name == "Outro":
        erp_name = st.text_input("Digite o nome do ERP", value="Meu ERP")

with col2:
    st.subheader("📁 Arquivo 2: Vendas Operadora")
    file_operadora = st.file_uploader("Upload Planilha Operadora", type=['csv', 'xlsx', 'xls'], key="operadora")
    operadora_source = st.selectbox("Fonte dos Dados", options=["Sistema Conciliadora", "Outros"])
    
    if operadora_source == "Outros":
        operadora_source = st.text_input("Digite a Fonte dos Dados", value="Portal Operadora")

# --- Lógica de Processamento ---

# --- Inicialização do Session State ---
if 'step' not in st.session_state:
    st.session_state.step = 0
if 'df_erp_processed' not in st.session_state:
    st.session_state.df_erp_processed = None
if 'df_op_processed' not in st.session_state:
    st.session_state.df_op_processed = None

# --- Lógica de Processamento ---

if file_erp and file_operadora:
    df_erp = load_data(file_erp)
    df_operadora = load_data(file_operadora)

    if df_erp is not None and df_operadora is not None:
        st.divider()
        
        # --- ETAPA 0: Configuração e Processamento Inicial ---
        if st.session_state.step == 0:
            st.header("🛠️ Mapeamento de Colunas")
            
            # Colunas Padrão Requeridas (ERP)
            # Colunas Padrão Requeridas (ERP)
            required_columns_erp = {
                "data_venda": "Data da Venda",
                "valor_venda": "Valor da Venda",
                "operadora": "Operadora",
                "nsu": "NSU",
                "parcelas": "Parcelas",
                "taxa": "Taxa (%)",
                "valor_liquido": "Valor Líquido",
                "valor_taxa": "Valor da Taxa"
            }

            # Ajuste dinâmico para Brajan
            if erp_name == "Brajan":
                if "taxa" in required_columns_erp:
                    del required_columns_erp["taxa"]
                if "valor_liquido" in required_columns_erp:
                    del required_columns_erp["valor_liquido"]

            # Colunas Padrão Requeridas (Operadora) - Sem Parcelas
            required_columns_op = {
                "data_venda": "Data da Venda",
                "valor_venda": "Valor da Venda",
                "operadora": "Operadora",
                "nsu": "NSU",
                "taxa": "Taxa (%)",
                "valor_liquido": "Valor Líquido",
                "valor_pago": "Valor Pago",
                "taxa_pagamento": "Taxa Pagamento (%)",
                "valor_taxa": "Valor da Taxa"
            }

            # Ajuste dinâmico para Sistema Conciliadora
            if operadora_source == "Sistema Conciliadora":
                if "valor_pago" in required_columns_op:
                    del required_columns_op["valor_pago"]

            # Mapeamentos Automáticos (Padrão)
            ERP_MAPPINGS = {
                "Brajan": {
                    "data_venda": "Data Efetiva",
                    "valor_venda": "Vr. Venda",
                    "operadora": "Bandeira",
                    "nsu": "Autorização",
                    "parcelas": "Pc",
                    "taxa": "Taxa",
                    "valor_liquido": "Liquido",
                    "valor_taxa": "Taxa"
                }
            }

            OPERADORA_MAPPINGS = {
                "Sistema Conciliadora": {
                    "data_venda": "Venda",
                    "valor_venda": "Bruto",
                    "operadora": "Produto",
                    "nsu": "NSU",
                    "taxa": "Taxa",
                    "valor_liquido": "Liquido",
                    "taxa_pagamento": "Taxa",
                    "valor_taxa": "Valor Taxa"
                }
            }

            col_map1, col_map2 = st.columns(2)

            mapping_erp = {}
            mapping_operadora = {}

            # Mapeamento ERP
            with col_map1:
                st.subheader(f"Mapear: {erp_name}")
                st.info("Selecione a coluna correspondente no seu arquivo.")
                
                # Obter mapeamento padrão se existir
                current_erp_map = ERP_MAPPINGS.get(erp_name, {})
                
                for key, label in required_columns_erp.items():
                    # Tentar encontrar o índice da coluna padrão
                    default_col_name = current_erp_map.get(key)
                    default_index = 0
                    
                    if default_col_name and default_col_name in df_erp.columns:
                        default_index = list(df_erp.columns).index(default_col_name)
                    
                    mapping_erp[key] = st.selectbox(
                        f"{label} ({erp_name})",
                        options=df_erp.columns,
                        index=default_index,
                        key=f"erp_{key}"
                    )

            # Mapeamento Operadora
            with col_map2:
                st.subheader(f"Mapear: {operadora_source}")
                st.info("Selecione a coluna correspondente no seu arquivo.")
                
                # Obter mapeamento padrão se existir
                current_op_map = OPERADORA_MAPPINGS.get(operadora_source, {})

                for key, label in required_columns_op.items():
                    # Tentar encontrar o índice da coluna padrão
                    default_col_name = current_op_map.get(key)
                    default_index = 0
                    
                    if default_col_name and default_col_name in df_operadora.columns:
                        default_index = list(df_operadora.columns).index(default_col_name)

                    mapping_operadora[key] = st.selectbox(
                        f"{label} ({operadora_source})",
                        options=df_operadora.columns,
                        index=default_index,
                        key=f"op_{key}"
                    )

            st.divider()
            st.divider()
            st.header("⚙️ Opções de Conciliação")
            
            col_opt1, col_opt2 = st.columns(2)
            with col_opt1:
                check_erp_filter = st.checkbox("Filtrar Vendas Parceladas (Apenas Parcelas = 1)", value=True)
                check_date_fmt = st.checkbox("Padronizar formato de datas (DD/MM/AAAA)", value=True)
                check_debits = st.checkbox("Separar débitos (- R$)", value=True)
            with col_opt2:
                check_op_agg = st.checkbox("Unificar Vendas por NSU (Somar Valores)", value=True)
                check_val_fmt = st.checkbox("Padronizar formato de valores (R$ 0.00)", value=True)
            
            check_depara = st.checkbox("Padronizar Operadoras (De-Para)", value=True)

            st.divider()
            st.header("📅 Filtros e Processamento")

            # Filtro de Data
            col_date1, col_date2 = st.columns(2)
            with col_date1:
                start_date = st.date_input("Data Início")
            with col_date2:
                end_date = st.date_input("Data Fim")

            if st.button("🚀 Processar Dados", type="primary"):
                try:
                    # Inverter o dicionário para usar no rename (De: Original -> Para: Padrão)
                    rename_map_erp = {v: k for k, v in mapping_erp.items()}
                    rename_map_op = {v: k for k, v in mapping_operadora.items()}

                    # Padronização
                    df_erp_std = standardize_df(df_erp, rename_map_erp, "origem", erp_name)
                    df_op_std = standardize_df(df_operadora, rename_map_op, "origem", operadora_source)

                    # Conversão de Data (Necessário para Agregação e Filtragem)
                    df_erp_std['data_venda'] = pd.to_datetime(df_erp_std['data_venda'], errors='coerce')
                    df_op_std['data_venda'] = pd.to_datetime(df_op_std['data_venda'], errors='coerce')

                    # --- Regras de Negócio ---
                    
                    # 1. ERP: Filtrar Parcelas == '1'
                    if check_erp_filter:
                        # Converter para string e remover decimais (.0) caso existam
                        if 'parcelas' in df_erp_std.columns:
                            df_erp_std['parcelas'] = df_erp_std['parcelas'].astype(str).str.replace(r'\.0$', '', regex=True)
                            df_erp_std = df_erp_std[df_erp_std['parcelas'] == '1']

                    # 2. Separar Débitos (Valores Negativos)
                    df_debits_final = pd.DataFrame()
                    if check_debits:
                        # Converter para numérico para filtrar
                        df_erp_std['valor_venda_temp'] = pd.to_numeric(df_erp_std['valor_venda'], errors='coerce').fillna(0)
                        df_op_std['valor_venda_temp'] = pd.to_numeric(df_op_std['valor_venda'], errors='coerce').fillna(0)

                        # Filtrar Débitos
                        df_erp_debits = df_erp_std[df_erp_std['valor_venda_temp'] < 0].copy()
                        df_op_debits = df_op_std[df_op_std['valor_venda_temp'] < 0].copy()

                        # Adicionar coluna Origem
                        df_erp_debits['Origem'] = 'ERP'
                        df_op_debits['Origem'] = 'Operadora'

                        # Concatenar
                        df_debits_final = pd.concat([df_erp_debits, df_op_debits], ignore_index=True)
                        
                        # Remover colunas temporárias e desnecessárias dos débitos
                        if 'valor_venda_temp' in df_debits_final.columns:
                            df_debits_final = df_debits_final.drop(columns=['valor_venda_temp'])

                        # Remover Débitos dos DataFrames principais
                        df_erp_std = df_erp_std[df_erp_std['valor_venda_temp'] >= 0].drop(columns=['valor_venda_temp'])
                        df_op_std = df_op_std[df_op_std['valor_venda_temp'] >= 0].drop(columns=['valor_venda_temp'])
                    
                    # Salvar Débitos no Session State
                    st.session_state.df_debits = df_debits_final

                    # 3. Operadora: Unificar por NSU (Somar Valor)
                    df_op_std['valor_venda'] = pd.to_numeric(df_op_std['valor_venda'], errors='coerce').fillna(0)
                    
                    # Converter novas colunas para numérico se existirem (Operadora)
                    for col in ['taxa', 'valor_liquido', 'valor_pago', 'taxa_pagamento', 'valor_taxa']:
                        if col in df_op_std.columns:
                            df_op_std[col] = df_op_std[col].apply(clean_input_value)

                    # Converter novas colunas para numérico se existirem (ERP)
                    for col in ['taxa', 'valor_liquido', 'valor_taxa']:
                        if col in df_erp_std.columns:
                            df_erp_std[col] = df_erp_std[col].apply(clean_input_value)
                    
                    if check_op_agg:
                        if 'nsu' in df_op_std.columns:
                            # Definir dicionário de agregação
                            agg_dict = {
                                'valor_venda': 'sum',
                                'origem': 'first'
                            }
                            # Adicionar novas colunas ao dicionário de agregação se existirem
                            for col in ['taxa', 'valor_liquido', 'valor_pago', 'taxa_pagamento', 'valor_taxa']:
                                if col in df_op_std.columns:
                                    if 'taxa' in col and col != 'valor_taxa':
                                        agg_dict[col] = 'mean' # Média para taxas
                                    else:
                                        agg_dict[col] = 'sum' # Soma para valores
                            
                            # Agrupar por NSU, Data e Operadora
                            df_op_std = df_op_std.groupby(['nsu', 'data_venda', 'operadora'], as_index=False).agg(agg_dict)

                    # Filtragem por Data
                    start_date = pd.to_datetime(start_date)
                    end_date = pd.to_datetime(end_date)

                    mask_erp = (df_erp_std['data_venda'] >= start_date) & (df_erp_std['data_venda'] <= end_date)
                    mask_op = (df_op_std['data_venda'] >= start_date) & (df_op_std['data_venda'] <= end_date)

                    df_erp_final = df_erp_std.loc[mask_erp].copy()
                    df_op_final = df_op_std.loc[mask_op].copy()

                    # --- Formatação de Dados (Opcional) ---
                    # Aplicar formatação APÓS filtragem e ANTES de salvar no session state
                    # Isso converte para string, o que é bom para visualização e merge exato
                    
                    if check_date_fmt:
                        df_erp_final['data_venda'] = df_erp_final['data_venda'].dt.strftime('%d/%m/%Y')
                        df_op_final['data_venda'] = df_op_final['data_venda'].dt.strftime('%d/%m/%Y')
                    
                    if check_val_fmt:
                        # Função auxiliar para formatar moeda
                        def format_currency(val):
                            return f"R$ {val:.2f}"
                        
                        # Garantir que é numérico antes de formatar
                        df_erp_final['valor_venda'] = pd.to_numeric(df_erp_final['valor_venda'], errors='coerce').fillna(0)
                        df_op_final['valor_venda'] = pd.to_numeric(df_op_final['valor_venda'], errors='coerce').fillna(0)
                        
                        df_erp_final['valor_venda'] = df_erp_final['valor_venda'].apply(format_currency)
                        df_op_final['valor_venda'] = df_op_final['valor_venda'].apply(format_currency)
                        
                        # Formatar também valor_taxa se existir
                        if 'valor_taxa' in df_erp_final.columns:
                            df_erp_final['valor_taxa'] = pd.to_numeric(df_erp_final['valor_taxa'], errors='coerce').fillna(0)
                            df_erp_final['valor_taxa'] = df_erp_final['valor_taxa'].apply(format_currency)
                        if 'valor_taxa' in df_op_final.columns:
                            df_op_final['valor_taxa'] = pd.to_numeric(df_op_final['valor_taxa'], errors='coerce').fillna(0)
                            df_op_final['valor_taxa'] = df_op_final['valor_taxa'].apply(format_currency)

                    # Salvar no Session State e avançar
                    st.session_state.df_erp_processed = df_erp_final
                    st.session_state.df_op_processed = df_op_final
                    st.session_state.start_date = start_date
                    st.session_state.end_date = end_date
                    
                    # Lógica de Pulo do De-Para
                    if check_depara:
                        st.session_state.step = 1
                    else:
                        st.session_state.step = 2
                    
                    st.rerun()

                except Exception as e:
                    st.error(f"Ocorreu um erro durante o processamento: {e}")
                    st.exception(e)

        # --- ETAPA 1: De-Para de Operadoras ---
        elif st.session_state.step == 1:
            st.header("🔄 Padronização de Operadoras (De-Para)")
            st.info("Identifique e padronize os nomes das operadoras encontradas. Deixe em branco para manter o original.")

            df_erp_final = st.session_state.df_erp_processed
            df_op_final = st.session_state.df_op_processed

            # Obter valores únicos
            unique_ops_erp = sorted(df_erp_final['operadora'].astype(str).unique())
            unique_ops_op = sorted(df_op_final['operadora'].astype(str).unique())
            
            # Combinar valores únicos de ambas as planilhas para facilitar
            all_unique_ops = sorted(list(set(unique_ops_erp + unique_ops_op)))

            st.write(f"Foram encontradas {len(all_unique_ops)} variações de operadoras.")

            # Formulário para De-Para
            with st.form("de_para_form"):
                col_depara1, col_depara2 = st.columns(2)
                
                mapping_changes = {}
                
                # Distribuir inputs em duas colunas
                # Distribuir inputs em duas colunas
                for i, op_name in enumerate(all_unique_ops):
                    col = col_depara1 if i % 2 == 0 else col_depara2
                    
                    # Obter sugestão automática
                    suggestion = get_operator_suggestion(op_name)
                    
                    # Criar colunas para checkbox e input
                    c1, c2 = col.columns([0.15, 0.85])
                    
                    # Checkbox para incluir/excluir
                    include_op = c1.checkbox("Incluir", value=True, key=f"check_{i}", label_visibility="collapsed")
                    
                    if not include_op:
                        mapping_changes[op_name] = "__REMOVE__"
                    
                    # Input para renomear
                    new_val = c2.text_input(f"De: '{op_name}' Para:", value=suggestion, key=f"depara_{i}")
                    
                    if include_op and new_val.strip():
                        mapping_changes[op_name] = new_val.strip()

                submitted = st.form_submit_button("✅ Finalizar e Visualizar Resultados")
                
                if submitted:
                    # Identificar operadoras a remover
                    ops_to_remove = [k for k, v in mapping_changes.items() if v == "__REMOVE__"]
                    
                    # Filtrar DataFrames
                    if ops_to_remove:
                        df_erp_final = df_erp_final[~df_erp_final['operadora'].astype(str).isin(ops_to_remove)]
                        df_op_final = df_op_final[~df_op_final['operadora'].astype(str).isin(ops_to_remove)]
                        st.warning(f"{len(ops_to_remove)} operadoras foram removidas da análise.")

                    # Aplicar mudanças de nome (apenas para os que não foram removidos)
                    rename_map = {k: v for k, v in mapping_changes.items() if v != "__REMOVE__"}
                    
                    if rename_map:
                        df_erp_final['operadora'] = df_erp_final['operadora'].astype(str).replace(rename_map)
                        df_op_final['operadora'] = df_op_final['operadora'].astype(str).replace(rename_map)
                        st.success(f"{len(rename_map)} padronizações aplicadas!")
                    
                    # Atualizar Session State
                    st.session_state.df_erp_processed = df_erp_final
                    st.session_state.df_op_processed = df_op_final
                    st.session_state.step = 2
                    st.rerun()

        # --- ETAPA 2: Resultados da Conciliação ---
        elif st.session_state.step == 2:
            st.header("📊 Resultados da Conciliação")
            st.info("Avaliação considerando as variáveis: 'Data da Venda', 'Valor da Venda' e 'Operadora'")

            df_erp = st.session_state.df_erp_processed.copy()
            df_op = st.session_state.df_op_processed.copy()

            # --- Lógica de Conciliação 1-para-1 ---
            
            # Definir colunas chave para conciliação
            match_cols = ['data_venda', 'valor_venda', 'operadora']
            
            # Criar ID único para tratar duplicatas (cumcount)
            # Isso garante que se houver 2 vendas iguais no ERP e 2 na Operadora, elas sejam casadas 1-a-1
            df_erp['match_id'] = df_erp.groupby(match_cols).cumcount()
            df_op['match_id'] = df_op.groupby(match_cols).cumcount()

            # Realizar o Merge (Left Join: ERP -> Operadora)
            # Indicador=True cria a coluna '_merge' que diz se encontrou ou não
            df_merged = pd.merge(
                df_erp,
                df_op[match_cols + ['match_id']], # Trazer apenas as chaves para verificar existência
                on=match_cols + ['match_id'],
                how='left',
                indicator=True
            )

            # Definir Status
            # both = Encontrado em ambos (CONCILIADO)
            # left_only = Encontrado apenas no ERP (DIVERGENTE)
            df_merged['status_conciliacao'] = df_merged['_merge'].map({
                'both': 'CONCILIADO',
                'left_only': 'DIVERGENTE'
            })

            # Remover colunas auxiliares
            df_merged = df_merged.drop(columns=['match_id', '_merge'])

            # Métricas ERP
            total_vendas = len(df_merged)
            total_conciliado = len(df_merged[df_merged['status_conciliacao'] == 'CONCILIADO'])
            total_divergente = len(df_merged[df_merged['status_conciliacao'] == 'DIVERGENTE'])

            # --- Conciliação Inversa (Operadora -> ERP) ---
            # Realizar o Merge (Left Join: Operadora -> ERP)
            df_merged_op = pd.merge(
                df_op,
                df_erp[match_cols + ['match_id']], 
                on=match_cols + ['match_id'],
                how='left',
                indicator=True
            )

            df_merged_op['status_conciliacao'] = df_merged_op['_merge'].map({
                'both': 'CONCILIADO',
                'left_only': 'DIVERGENTE'
            })
            df_merged_op = df_merged_op.drop(columns=['match_id', '_merge'])

            # Métricas Operadora
            total_vendas_op = len(df_merged_op)
            total_conciliado_op = len(df_merged_op[df_merged_op['status_conciliacao'] == 'CONCILIADO'])
            total_divergente_op = len(df_merged_op[df_merged_op['status_conciliacao'] == 'DIVERGENTE'])

            # Exibir Métricas (ERP)
            st.markdown("### Métricas ERP")
            col_met1, col_met2, col_met3 = st.columns(3)
            col_met1.metric("Total de Vendas (ERP)", total_vendas)
            col_met2.metric("Conciliadas (ERP)", total_conciliado)
            col_met3.metric("Divergentes (ERP)", total_divergente)
            
            # Exibir Métricas (Operadora)
            st.markdown("### Métricas Operadora")
            col_met_op1, col_met_op2, col_met_op3 = st.columns(3)
            col_met_op1.metric("Total de Vendas (Operadora)", total_vendas_op)
            col_met_op2.metric("Conciliadas (Operadora)", total_conciliado_op)
            col_met_op3.metric("Divergentes (Operadora)", total_divergente_op)

            st.divider()
            st.subheader("📊 Resumo dos Resultados")
            
            # --- Cálculo de Métricas Financeiras ---
            

            
            # --- PREPARAÇÃO DOS DADOS (OPERADORA) ---
            df_op_metrics = df_op.copy()
            metric_cols_op = ['valor_venda', 'taxa', 'valor_liquido', 'valor_taxa']
            
            for col in metric_cols_op:
                if col in df_op_metrics.columns:
                     df_op_metrics[col] = df_op_metrics[col].apply(clean_input_value)
                else:
                    df_op_metrics[col] = 0.0

            # --- PREPARAÇÃO DOS DADOS (ERP) ---
            df_erp_metrics = df_erp.copy()
            metric_cols_erp = ['valor_venda', 'taxa', 'valor_liquido', 'valor_taxa']

            for col in metric_cols_erp:
                if col in df_erp_metrics.columns:
                     df_erp_metrics[col] = df_erp_metrics[col].apply(clean_input_value)
                else:
                    df_erp_metrics[col] = 0.0

            # --- CÁLCULO DE MÉTRICAS (OPERADORA) ---
            # 1. Valor total de vendas brutas
            vendas_brutas_op = df_op_metrics['valor_venda'].sum()
            
            # 2. Taxa média aplicada nas vendas
            # 2. Valor total da taxa aplicada nas vendas
            if df_op_metrics['valor_taxa'].sum() > 0:
                valor_total_taxa_op = df_op_metrics['valor_taxa'].sum()
            elif df_op_metrics['valor_liquido'].sum() > 0:
                 valor_total_taxa_op = vendas_brutas_op - df_op_metrics['valor_liquido'].sum()
            else:
                 valor_total_taxa_op = (df_op_metrics['valor_venda'] * df_op_metrics['taxa'] / 100).sum()

            # 3. Taxa média aplicada nas vendas
            if vendas_brutas_op > 0:
                taxa_media_op = (valor_total_taxa_op / vendas_brutas_op) * 100
            else:
                taxa_media_op = 0.0

            # Vendas Líquidas
            if df_op_metrics['valor_liquido'].sum() > 0:
                 vendas_liquidas_op = df_op_metrics['valor_liquido'].sum()
            else:
                 vendas_liquidas_op = vendas_brutas_op - valor_total_taxa_op

            # 4. Débitos em vendas
            debitos_op_val = 0.0
            if 'df_debits' in st.session_state and st.session_state.df_debits is not None and not st.session_state.df_debits.empty:
                 debitos_op_df = st.session_state.df_debits[st.session_state.df_debits['Origem'] == 'Operadora']
                 if not debitos_op_df.empty:
                     debitos_op_val = pd.to_numeric(debitos_op_df['valor_venda'], errors='coerce').sum()

            # --- CÁLCULO DE MÉTRICAS (ERP) ---
            # 1. Valor total de vendas brutas
            vendas_brutas_erp = df_erp_metrics['valor_venda'].sum()
            
            # 2. Taxa média aplicada nas vendas
            # 2. Valor total da taxa aplicada nas vendas
            if df_erp_metrics['valor_taxa'].sum() > 0:
                valor_total_taxa_erp = df_erp_metrics['valor_taxa'].sum()
            elif df_erp_metrics['valor_liquido'].sum() > 0:
                 valor_total_taxa_erp = vendas_brutas_erp - df_erp_metrics['valor_liquido'].sum()
            else:
                 valor_total_taxa_erp = (df_erp_metrics['valor_venda'] * df_erp_metrics['taxa'] / 100).sum()

            # 3. Taxa média aplicada nas vendas
            if vendas_brutas_erp > 0:
                taxa_media_erp = (valor_total_taxa_erp / vendas_brutas_erp) * 100
            else:
                taxa_media_erp = 0.0

            # Vendas Líquidas
            if df_erp_metrics['valor_liquido'].sum() > 0:
                 vendas_liquidas_erp = df_erp_metrics['valor_liquido'].sum()
            else:
                 vendas_liquidas_erp = vendas_brutas_erp - valor_total_taxa_erp

            # 4. Débitos em vendas
            debitos_erp_val = 0.0
            if 'df_debits' in st.session_state and st.session_state.df_debits is not None and not st.session_state.df_debits.empty:
                 debitos_erp_df = st.session_state.df_debits[st.session_state.df_debits['Origem'] == 'ERP']
                 if not debitos_erp_df.empty:
                     debitos_erp_val = pd.to_numeric(debitos_erp_df['valor_venda'], errors='coerce').sum()

            # Exibir Tabela de Resumo Comparativo
            st.markdown("#### Detalhamento Financeiro")
            
            resumo_data = {
                "Métrica": [
                    "Valor total de vendas brutas",
                    "Taxa média aplicada nas vendas",
                    "Valor total da taxa aplicada nas vendas",
                    "Débitos em vendas",
                    "Valor total de vendas líquidas"
                ],
                "Valor - ERP": [
                    f"R$ {vendas_brutas_erp:,.2f}",
                    f"{taxa_media_erp:.2f}%",
                    f"R$ {valor_total_taxa_erp:,.2f}",
                    f"R$ {debitos_erp_val:,.2f}",
                    f"R$ {vendas_liquidas_erp:,.2f}"
                ],
                "Valor - Operadora": [
                    f"R$ {vendas_brutas_op:,.2f}",
                    f"{taxa_media_op:.2f}%",
                    f"R$ {valor_total_taxa_op:,.2f}",
                    f"R$ {debitos_op_val:,.2f}",
                    f"R$ {vendas_liquidas_op:,.2f}"
                ]
            }
            
            df_resumo_financeiro = pd.DataFrame(resumo_data)
            st.dataframe(df_resumo_financeiro, width='stretch', hide_index=True)

            st.divider()
            st.header("📈 Resultados Gerais")
            
            st.markdown("#### Resultado por adquirente - No período")

            # --- Cálculo do Resumo por Operadora ---
            # Precisamos garantir que os valores sejam numéricos para somar
            # Se a formatação foi aplicada, 'valor_venda' será string (R$ ...). Precisamos limpar.
            
            def clean_currency(val):
                if isinstance(val, str):
                    # Remover 'R$', espaços e substituir vírgula por ponto se necessário (embora nosso format use ponto)
                    clean_val = val.replace('R$', '').replace(' ', '')
                    return float(clean_val)
                return float(val)

            df_erp_calc = df_erp.copy()
            df_op_calc = df_op.copy()

            df_erp_calc['valor_venda_num'] = df_erp_calc['valor_venda'].apply(clean_currency)
            df_op_calc['valor_venda_num'] = df_op_calc['valor_venda'].apply(clean_currency)

            # Agrupamento
            erp_summary = df_erp_calc.groupby('operadora').agg(
                valor_venda_num=('valor_venda_num', 'sum'),
                qtd_vendas_erp=('operadora', 'count')
            ).reset_index()
            
            op_summary = df_op_calc.groupby('operadora').agg(
                valor_venda_num=('valor_venda_num', 'sum'),
                qtd_vendas_op=('operadora', 'count')
            ).reset_index()

            # Merge dos resumos
            summary_merged = pd.merge(erp_summary, op_summary, on='operadora', how='outer', suffixes=('_erp', '_op')).fillna(0)

            # Cálculo da Divergência
            summary_merged['divergencia'] = summary_merged['valor_venda_num_erp'] - summary_merged['valor_venda_num_op']
            summary_merged['divergencia_qtd'] = summary_merged['qtd_vendas_erp'] - summary_merged['qtd_vendas_op']

            # Renomear e Formatar para Exibição
            summary_display = summary_merged.rename(columns={
                'operadora': 'Operadora',
                'valor_venda_num_erp': 'Valor de venda total no ERP',
                'valor_venda_num_op': 'Valor de venda total na Operadora',
                'divergencia': 'Divergência de valor',
                'qtd_vendas_erp': 'Quantidade de vendas no ERP',
                'qtd_vendas_op': 'Quantidade de vendas na Operadora'
            })

            # Aplicar formatação de moeda para exibição
            def fmt_currency(x): return f"R$ {x:.2f}"
            
            summary_display['Valor de venda total no ERP'] = summary_display['Valor de venda total no ERP'].apply(fmt_currency)
            summary_display['Valor de venda total na Operadora'] = summary_display['Valor de venda total na Operadora'].apply(fmt_currency)
            summary_display['Divergência de valor'] = summary_display['Divergência de valor'].apply(fmt_currency)
            
            # Formatar quantidades como inteiros
            summary_display['Quantidade de vendas no ERP'] = summary_display['Quantidade de vendas no ERP'].astype(int)
            summary_display['Quantidade de vendas na Operadora'] = summary_display['Quantidade de vendas na Operadora'].astype(int)

            # Reordenar colunas
            cols_order = [
                'Operadora',
                'Valor de venda total no ERP',
                'Valor de venda total na Operadora',
                'Divergência de valor',
                'Quantidade de vendas no ERP',
                'Quantidade de vendas na Operadora'
            ]
            summary_display = summary_display[cols_order]

            # Filtros Dinâmicos (Resultado por adquirente - No período)
            with st.expander("Filtros Avançados (Resultado por adquirente - No período)", expanded=False):
                st.info("Selecione os valores para filtrar a tabela abaixo. Deixe em branco para ver tudo.")
                
                summary_filtered = summary_display.copy()
                
                # Criar colunas para os filtros (layout em grade)
                filter_cols_summary = st.columns(3)
                
                # Iterar sobre as colunas para criar filtros
                for i, col in enumerate(summary_display.columns):
                    unique_vals_summary = sorted(summary_display[col].astype(str).unique())
                    
                    with filter_cols_summary[i % 3]:
                        selected_vals_summary = st.multiselect(f"{col}", unique_vals_summary, key=f"filter_summary_{col}")
                        
                        if selected_vals_summary:
                            summary_filtered = summary_filtered[summary_filtered[col].astype(str).isin(selected_vals_summary)]

            st.dataframe(summary_filtered, width='stretch')

            st.divider()
            st.markdown("#### Resultado por adquirente - Diário")

            # --- Cálculo do Resumo Diário por Operadora ---
            # Agrupamento por Data e Operadora
            erp_daily = df_erp_calc.groupby(['data_venda', 'operadora']).agg(
                valor_venda_num=('valor_venda_num', 'sum'),
                qtd_vendas_erp=('operadora', 'count')
            ).reset_index()
            
            op_daily = df_op_calc.groupby(['data_venda', 'operadora']).agg(
                valor_venda_num=('valor_venda_num', 'sum'),
                qtd_vendas_op=('operadora', 'count')
            ).reset_index()

            # Merge dos resumos diários
            daily_merged = pd.merge(erp_daily, op_daily, on=['data_venda', 'operadora'], how='outer', suffixes=('_erp', '_op')).fillna(0)

            # Cálculo da Divergência
            daily_merged['divergencia'] = daily_merged['valor_venda_num_erp'] - daily_merged['valor_venda_num_op']
            daily_merged['divergencia_qtd'] = daily_merged['qtd_vendas_erp'] - daily_merged['qtd_vendas_op']

            # Renomear e Formatar para Exibição
            daily_display = daily_merged.rename(columns={
                'data_venda': 'Data',
                'operadora': 'Operadora',
                'valor_venda_num_erp': 'Valor de venda total no ERP',
                'valor_venda_num_op': 'Valor de venda total na Operadora',
                'divergencia': 'Divergência de valor',
                'qtd_vendas_erp': 'Quantidade de vendas no ERP',
                'qtd_vendas_op': 'Quantidade de vendas na Operadora'
            })

            # Selecionar colunas finais
            cols_to_show = ['Data', 'Operadora', 'Valor de venda total no ERP', 'Valor de venda total na Operadora', 'Divergência de valor', 'Quantidade de vendas no ERP', 'Quantidade de vendas na Operadora']
            daily_display = daily_display[cols_to_show]

            # Aplicar formatação de moeda
            daily_display['Valor de venda total no ERP'] = daily_display['Valor de venda total no ERP'].apply(fmt_currency)
            daily_display['Valor de venda total na Operadora'] = daily_display['Valor de venda total na Operadora'].apply(fmt_currency)
            daily_display['Divergência de valor'] = daily_display['Divergência de valor'].apply(fmt_currency)
            
            # Formatar quantidades como inteiros
            daily_display['Quantidade de vendas no ERP'] = daily_display['Quantidade de vendas no ERP'].astype(int)
            daily_display['Quantidade de vendas na Operadora'] = daily_display['Quantidade de vendas na Operadora'].astype(int)

            # Filtros Dinâmicos (Resultado por adquirente - Diário)
            with st.expander("Filtros Avançados (Resultado por adquirente - Diário)", expanded=False):
                st.info("Selecione os valores para filtrar a tabela abaixo. Deixe em branco para ver tudo.")
                
                daily_filtered = daily_display.copy()
                
                # Criar colunas para os filtros (layout em grade)
                filter_cols_daily = st.columns(3)
                
                # Iterar sobre as colunas para criar filtros
                for i, col in enumerate(daily_display.columns):
                    unique_vals_daily = sorted(daily_display[col].astype(str).unique())
                    
                    with filter_cols_daily[i % 3]:
                        selected_vals_daily = st.multiselect(f"{col}", unique_vals_daily, key=f"filter_daily_{col}")
                        
                        if selected_vals_daily:
                            daily_filtered = daily_filtered[daily_filtered[col].astype(str).isin(selected_vals_daily)]

            st.dataframe(daily_filtered, width='stretch')

            st.divider()
            st.header("🔍 Detalhamento - ERP -> Operadora")
            
            # --- Filtros Dinâmicos (Detalhamento ERP -> Operadora) ---
            with st.expander("Filtros Avançados (ERP -> Operadora)", expanded=False):
                st.info("Selecione os valores para filtrar a tabela abaixo. Deixe em branco para ver tudo.")
                
                df_filtered = df_merged.copy()
                
                # Criar colunas para os filtros (layout em grade)
                filter_cols = st.columns(3)
                
                # Iterar sobre as colunas para criar filtros
                for i, col in enumerate(df_merged.columns):
                    # Pular colunas que não fazem sentido filtrar ou são muito únicas (opcional)
                    if col == 'match_id': continue
                    
                    unique_vals = sorted(df_merged[col].astype(str).unique())
                    
                    with filter_cols[i % 3]:
                        selected_vals = st.multiselect(f"{col}", unique_vals, key=f"filter_erp_{col}")
                        
                        if selected_vals:
                            df_filtered = df_filtered[df_filtered[col].astype(str).isin(selected_vals)]

            st.dataframe(df_filtered, width='stretch')

            st.divider()
            st.header("🔍 Detalhamento - Operadora -> ERP")
            
            # --- Filtros Dinâmicos (Detalhamento Operadora -> ERP) ---
            with st.expander("Filtros Avançados (Operadora -> ERP)", expanded=False):
                st.info("Selecione os valores para filtrar a tabela abaixo. Deixe em branco para ver tudo.")
                
                df_filtered_op = df_merged_op.copy()
                
                # Criar colunas para os filtros (layout em grade)
                filter_cols_op = st.columns(3)
                
                # Iterar sobre as colunas para criar filtros
                for i, col in enumerate(df_merged_op.columns):
                    # Pular colunas que não fazem sentido filtrar ou são muito únicas (opcional)
                    if col == 'match_id': continue
                    
                    unique_vals_op = sorted(df_merged_op[col].astype(str).unique())
                    
                    with filter_cols_op[i % 3]:
                        selected_vals_op = st.multiselect(f"{col}", unique_vals_op, key=f"filter_op_{col}")
                        
                        if selected_vals_op:
                            df_filtered_op = df_filtered_op[df_filtered_op[col].astype(str).isin(selected_vals_op)]

            st.dataframe(df_filtered_op, width='stretch')

            # Exibir Débitos Separados se houver
            if 'df_debits' in st.session_state and st.session_state.df_debits is not None and not st.session_state.df_debits.empty:
                st.divider()
                with st.expander("💸 Débitos Separados (- R$)", expanded=False):
                    st.warning("Estes valores foram separados do processamento principal por serem negativos.")
                    st.dataframe(st.session_state.df_debits, width='stretch')

            # Botão para Reiniciar
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("🔄 Reiniciar Processo"):
                    st.session_state.step = 0
                    st.session_state.df_erp_processed = None
                    st.session_state.df_op_processed = None
                    st.session_state.df_debits = None
                    st.rerun()
            
            with col_btn2:
                # Preparar dados para o relatório
                erp_metrics_dict = {
                    'vendas_brutas': vendas_brutas_erp,
                    'taxa_media': taxa_media_erp,
                    'total_taxas': valor_total_taxa_erp,
                    'vendas_liquidas': vendas_liquidas_erp
                }
                op_metrics_dict = {
                    'vendas_brutas': vendas_brutas_op,
                    'taxa_media': taxa_media_op,
                    'total_taxas': valor_total_taxa_op,
                    'vendas_liquidas': vendas_liquidas_op
                }
                
                # Filtrar Divergências
                df_erp_div = df_merged[df_merged['status_conciliacao'] == 'DIVERGENTE'].copy()
                df_op_div = df_merged_op[df_merged_op['status_conciliacao'] == 'DIVERGENTE'].copy()

                # Gerar Relatório
                report_buffer = generate_word_report(
                    erp_metrics_dict, 
                    op_metrics_dict, 
                    summary_display, 
                    daily_display, 
                    st.session_state.start_date, 
                    st.session_state.end_date,
                    df_erp_div,
                    df_op_div
                )
                
                st.download_button(
                    label="📄 Gerar Relatório (.docx)",
                    data=report_buffer,
                    file_name="relatorio_conciliacao.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
